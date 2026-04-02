"""文档解析服务测试。"""

from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app.services import document_parser
from app.services.document_parser import DocumentParseError, parse_document


def test_parse_txt_success(tmp_path: Path) -> None:
    """应能正确解析 txt 文件。"""
    file_path = tmp_path / "note.txt"
    file_path.write_text("hello rag", encoding="utf-8")

    doc = parse_document(file_path)

    assert doc.doc_type == "txt"
    assert doc.content == "hello rag"
    assert doc.metadata["file_name"] == "note.txt"


def test_parse_md_success(tmp_path: Path) -> None:
    """应能正确解析 md 文件。"""
    file_path = tmp_path / "readme.md"
    file_path.write_text("# title", encoding="utf-8")

    doc = parse_document(file_path)

    assert doc.doc_type == "md"
    assert doc.content == "# title"
    assert doc.metadata["extension"] == ".md"


def test_parse_missing_file_raise_error(tmp_path: Path) -> None:
    """文件不存在时应抛出解析异常。"""
    missing_file = tmp_path / "missing.txt"

    with pytest.raises(DocumentParseError, match="文件不存在"):
        parse_document(missing_file)


def test_parse_pdf_success(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """PDF 解析流程应被正确触发。"""
    file_path = tmp_path / "report.pdf"
    file_path.write_bytes(b"%PDF-1.7")

    # 避免依赖真实 PDF 结构，直接 mock 内部读取函数。
    monkeypatch.setattr(document_parser, "_read_pdf_file", lambda _path: "pdf body")

    doc = parse_document(file_path)

    assert doc.doc_type == "pdf"
    assert doc.content == "pdf body"
    assert doc.metadata["extension"] == ".pdf"
    assert doc.metadata["encoding"] == "binary"


def test_parse_docx_success(tmp_path: Path) -> None:
    """应能正确解析 docx 文件。"""
    file_path = tmp_path / "meeting.docx"
    doc = DocxDocument()
    doc.add_paragraph("line 1")
    doc.add_paragraph("line 2")
    doc.save(str(file_path))

    parsed = parse_document(file_path)

    assert parsed.doc_type == "docx"
    assert parsed.content == "line 1\nline 2"
    assert parsed.metadata["extension"] == ".docx"
    assert parsed.metadata["encoding"] == "binary"


def test_parse_unsupported_file_raise_error(tmp_path: Path) -> None:
    """不支持扩展名时应报错。"""
    file_path = tmp_path / "report.xlsx"
    file_path.write_text("fake-xlsx-content", encoding="utf-8")

    with pytest.raises(DocumentParseError, match="暂不支持的文件类型"):
        parse_document(file_path)


def test_parse_empty_file_raise_error(tmp_path: Path) -> None:
    """空内容文件应报错。"""
    file_path = tmp_path / "empty.txt"
    file_path.write_text("   \n\t", encoding="utf-8")

    with pytest.raises(DocumentParseError, match="文件内容为空"):
        parse_document(file_path)
