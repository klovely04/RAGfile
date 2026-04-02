"""文档解析服务。

支持格式：
- `.txt`
- `.md`
- `.pdf`
- `.docx`
"""

from pathlib import Path

from app.models import Document

# 纯文本类文档扩展名（可按编码直接读取）。
SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md"}
# 当前服务支持的全部扩展名。
SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class DocumentParseError(ValueError):
    """文档解析失败时抛出的统一异常类型。"""


def parse_document(file_path: str | Path) -> Document:
    """解析单个文件并返回统一 `Document` 模型。

    Args:
        file_path: 文件路径（`str` 或 `Path`）。

    Returns:
        Document: 标准化后的文档对象。

    Raises:
        DocumentParseError: 文件不存在、格式不支持、内容为空、读取失败等。
    """
    path = Path(file_path)

    # 基础校验：路径必须存在。
    if not path.exists():
        raise DocumentParseError(f"文件不存在: {path}")

    # 基础校验：路径必须是文件而不是目录。
    if not path.is_file():
        raise DocumentParseError(f"目标不是文件: {path}")

    # 读取扩展名并统一为小写。
    suffix = path.suffix.lower()
    # 校验扩展名是否支持。
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError(
            f"暂不支持的文件类型: {suffix}。当前仅支持: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    # 按文件类型路由到不同读取函数。
    if suffix in SUPPORTED_TEXT_EXTENSIONS:
        text, encoding = _read_text_file(path)
    elif suffix == ".pdf":
        text = _read_pdf_file(path)
        encoding = "binary"
    else:
        text = _read_docx_file(path)
        encoding = "binary"

    # 解析后的文本若为空，视为无效文档。
    if not text.strip():
        raise DocumentParseError(f"文件内容为空: {path}")

    # 返回统一模型。
    return Document(
        content=text,
        source=str(path),
        doc_type=suffix.lstrip("."),
        metadata={
            "file_name": path.name,
            "extension": suffix,
            "encoding": encoding,
        },
    )


def _read_text_file(path: Path) -> tuple[str, str]:
    """按候选编码读取文本文件。

    Args:
        path: 待读取文件路径。

    Returns:
        tuple[str, str]: `(文本内容, 实际使用编码)`。
    """
    # 常见中文场景编码优先级。
    candidate_encodings = ("utf-8", "utf-8-sig", "gb18030")

    for encoding in candidate_encodings:
        try:
            return path.read_text(encoding=encoding), encoding
        except UnicodeDecodeError:
            # 该编码不匹配时继续尝试下一个。
            continue

    raise DocumentParseError(f"文件编码无法识别，读取失败: {path}")


def _read_pdf_file(path: Path) -> str:
    """读取 PDF 文本并合并所有页面。

    Args:
        path: PDF 文件路径。

    Returns:
        str: 合并后的页面文本。
    """
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # pragma: no cover
        raise DocumentParseError(f"PDF 读取失败: {path}") from exc

    pages_text = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(pages_text)


def _read_docx_file(path: Path) -> str:
    """读取 DOCX 段落文本并按换行拼接。

    Args:
        path: DOCX 文件路径。

    Returns:
        str: 合并后的段落文本。
    """
    from docx import Document as DocxDocument

    try:
        doc = DocxDocument(str(path))
    except Exception as exc:  # pragma: no cover
        raise DocumentParseError(f"DOCX 读取失败: {path}") from exc

    return "\n".join(p.text for p in doc.paragraphs)
